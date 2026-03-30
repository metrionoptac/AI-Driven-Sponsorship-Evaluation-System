"""
Generates DOCX files from OrgRecords using python-docx.
Simulates Word document sponsorship requests with letterhead-style formatting.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .org_database import OrgRecord
from .templates import generate_letter_text


def build_docx(org: OrgRecord, output_dir: str) -> tuple[str, bytes]:
    """Build a DOCX sponsorship request letter. Returns (filepath, raw_bytes)."""
    letter_text = generate_letter_text(org)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Letterhead: org name right-aligned
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run(org.org_name or "")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Address if available
    if org.contact_address:
        addr_para = doc.add_paragraph()
        addr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = addr_para.add_run(org.contact_address)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph("")  # spacer

    # Body paragraphs
    for line in letter_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        else:
            # Check if it looks like a greeting/closing (bold it)
            is_greeting = any(line.startswith(g) for g in [
                "Sehr geehrte", "Hallo", "Guten Tag", "Hi", "Dear", "Hello",
                "Mit freundlichen", "Viele Gruesse", "Herzliche", "LG",
                "Kind regards", "Best regards", "Cheers", "Thanks",
            ])
            para = doc.add_paragraph()
            run = para.add_run(line)
            if is_greeting:
                run.bold = True

    # Save
    safe_name = org.org_name.replace(" ", "_").replace("/", "_")[:40] if org.org_name else org.id
    filename = f"{org.id}_{safe_name}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    with open(filepath, "rb") as f:
        docx_bytes = f.read()

    return filepath, docx_bytes
