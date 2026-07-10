"""
Test corpus generator — 50 emails + 50 web forms + 10 operator scenarios.

All data is FICTIONAL. Every sample gets a ground-truth entry in manifest.json
so the corpus doubles as the future eval-harness dataset (LLMOps track).

Region strategy (B30): the system is configured as Stadtwerke Bodensee, so the
corpus is majority Bodensee-region; ~15% out-of-region samples exist ON PURPOSE
to test region handling (B29). Umlauts written as ue/ae/oe (Windows console).

Run:  python test_data/generate_corpus.py
"""

import json
import os
import random
import sys
from email.message import EmailMessage
from email.utils import formatdate

import fitz  # PyMuPDF

try:
    import docx  # python-docx
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

random.seed(42)

ROOT = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(ROOT, "corpus")
DIR_EMAILS = os.path.join(OUT, "emails")
DIR_WEBFORM = os.path.join(OUT, "webform")
DIR_WF_ATT = os.path.join(OUT, "webform", "attachments")
DIR_OPERATOR = os.path.join(OUT, "operator")
for d in (DIR_EMAILS, DIR_WEBFORM, DIR_WF_ATT, DIR_OPERATOR):
    os.makedirs(d, exist_ok=True)

MANIFEST = []

# ----------------------------------------------------------------
# Data pools (fictional but plausible)
# ----------------------------------------------------------------

IN_REGION_TOWNS = [
    "Konstanz", "Friedrichshafen", "Lindau", "Singen", "Radolfzell",
    "Ueberlingen", "Meersburg", "Markdorf", "Stockach", "Salem",
    "Tettnang", "Kressbronn", "Sipplingen", "Bodman-Ludwigshafen",
]
OUT_REGION_TOWNS = ["Duesseldorf", "Berlin", "Muenchen", "Hamburg", "Leipzig"]

CLUB_TYPES = [
    ("SV",                     "sports",    "Jugendfussballturnier"),
    ("TSV",                    "sports",    "Anschaffung von Jugend-Trikots"),
    ("TV",                     "sports",    "Turnfest der Jugendabteilung"),
    ("FC",                     "sports",    "Fussballcamp fuer Kinder"),
    ("Musikverein",            "culture",   "Sommerkonzert im Stadtpark"),
    ("Kulturverein",           "culture",   "Kulturwochenende mit lokalen Kuenstlern"),
    ("Schachclub",             "sports",    "Jugend-Schachturnier"),
    ("DLRG Ortsgruppe",        "social",    "Schwimmkurse fuer Grundschulkinder"),
    ("Foerderverein Grundschule", "education", "Neue Buecher fuer die Schulbibliothek"),
    ("Tierschutzverein",       "social",    "Sanierung des Tierheims"),
    ("Jugendzentrum",          "social",    "Ferienprogramm fuer Jugendliche"),
    ("Naturfreunde",           "environment", "Baumpflanzaktion am Seeufer"),
]

FIRST = ["Anna", "Max", "Lena", "Jonas", "Marie", "Paul", "Laura", "Felix",
         "Sofia", "Lukas", "Emma", "David", "Julia", "Simon", "Clara", "Tim"]
LAST = ["Weber", "Schmid", "Keller", "Huber", "Maier", "Wagner", "Becker",
        "Braun", "Vogel", "Frank", "Berger", "Koch", "Richter", "Wolf"]

FREEMAIL = ["gmail.com", "gmx.de", "web.de", "t-online.de"]


def mk_org(i, town=None, club=None):
    club = club or CLUB_TYPES[i % len(CLUB_TYPES)]
    town = town or IN_REGION_TOWNS[i % len(IN_REGION_TOWNS)]
    prefix, category, purpose = club
    year = random.choice([1898, 1905, 1912, 1920, 1926, 1948, 1954, 1972, 1990])
    name = f"{prefix} {town} {year} e.V." if prefix in ("SV", "TSV", "TV", "FC") \
        else f"{prefix} {town} e.V."
    return {"name": name, "town": town, "category": category, "purpose": purpose}


def mk_contact(i, org_name, freemail=False):
    fn, ln = FIRST[i % len(FIRST)], LAST[(i * 3) % len(LAST)]
    if freemail:
        domain = random.choice(FREEMAIL)
        email = f"{fn.lower()}.{ln.lower()}{random.randint(1, 99)}@{domain}"
    else:
        slug = org_name.lower().replace(" ", "-").replace(".", "").replace("ev", "").strip("-")
        email = f"{fn.lower()}.{ln.lower()}@{slug[:40]}.de"
    return {"name": f"{fn} {ln}", "email": email,
            "phone": f"+49 75{random.randint(10, 99)} {random.randint(100000, 999999)}"}


def future_date():
    return f"2026-{random.randint(9, 12):02d}-{random.randint(1, 28):02d}"


def letter_lines(org, contact, amount, date, missing=()):
    """German request letter; omit fields listed in `missing`."""
    lines = [f"Sponsoringanfrage - {org['name']}", ""]
    lines.append(f"Verein: {org['name']}")
    if "contact" not in missing:
        lines.append(f"Ansprechpartner: {contact['name']}")
        lines.append(f"E-Mail: {contact['email']}")
        lines.append(f"Telefon: {contact['phone']}")
    lines.append(f"Region: {org['town']}")
    lines += ["", "Sehr geehrte Damen und Herren,", ""]
    ask = "wir bitten um Ihre Unterstuetzung"
    if "amount" not in missing:
        ask += f" in Hoehe von {amount:,.0f} EUR".replace(",", ".")
    ask += f" fuer folgendes Projekt: {org['purpose']}"
    if "date" not in missing:
        ask += f" am {date}"
    ask += f" in {org['town']}."
    lines.append(ask)
    lines.append(f"Erwartete Teilnehmer: ca. {random.choice([100, 200, 300, 400, 600, 800])}. "
                 f"Unser Verein zaehlt {random.choice([60, 85, 120, 200, 350])} Mitglieder.")
    lines += ["", "Als Gegenleistung bieten wir Logo-Praesenz auf Plakaten,",
              "Bannern und im Programmheft.", "",
              "Mit freundlichen Gruessen", contact["name"]]
    return lines


def write_pdf(path, lines):
    doc = fitz.open()
    page = doc.new_page()
    y = 72
    for i, line in enumerate(lines):
        page.insert_text((72, y), line, fontsize=15 if i == 0 else 11, fontname="helv")
        y += 26 if i == 0 else 17
    doc.save(path)
    doc.close()


def write_docx(path, lines):
    if not HAVE_DOCX:
        write_pdf(path.replace(".docx", ".pdf"), lines)
        return path.replace(".docx", ".pdf")
    d = docx.Document()
    d.add_heading(lines[0], level=1)
    for line in lines[1:]:
        d.add_paragraph(line)
    d.save(path)
    return path


def write_scan_jpg(path, lines):
    """Render the letter as an image -> simulates a scanned document (OCR path)."""
    tmp = path + ".tmp.pdf"
    write_pdf(tmp, lines)
    doc = fitz.open(tmp)
    pix = doc[0].get_pixmap(dpi=150)
    pix.save(path, output="jpeg")
    doc.close()
    os.remove(tmp)


def write_eml(path, sender, subject, body, attachments):
    """attachments: list of (filename, bytes, maintype, subtype)"""
    msg = EmailMessage()
    msg["From"] = sender
    msg["To"] = "sponsoring@stadtwerke-bodensee-demo.de"
    msg["Subject"] = subject
    msg["Date"] = formatdate(localtime=True)
    msg.set_content(body)
    for fname, data, maintype, subtype in attachments:
        msg.add_attachment(data, maintype=maintype, subtype=subtype, filename=fname)
    with open(path, "wb") as f:
        f.write(msg.as_bytes())


def email_body(org, contact, short=True):
    return (f"Sehr geehrtes Sponsoring-Team,\n\n"
            f"anbei erhalten Sie unsere Sponsoringanfrage ({org['purpose']}).\n"
            f"Bei Rueckfragen erreichen Sie mich unter {contact['phone']}.\n\n"
            f"Mit freundlichen Gruessen\n{contact['name']}\n{org['name']}")


def record(sid, channel, files, ground_truth, payload=None):
    entry = {"id": sid, "channel": channel, "files": files, "ground_truth": ground_truth}
    if payload is not None:
        entry["payload"] = payload
    MANIFEST.append(entry)


def gt(org, contact, amount, date, *, in_region=True, complete=True, missing=(),
       is_sponsorship=True, expect_eligible=True, needs_followup=False, notes=""):
    return {
        "organization": org["name"] if org else None,
        "amount": amount, "event_date": date,
        "region": org["town"] if org else None, "in_region": in_region,
        "category": org["category"] if org else None,
        "contact_email": contact["email"] if contact else None,
        "complete": complete, "missing_fields": list(missing),
        "is_sponsorship": is_sponsorship,
        "expect": {"eligible_hard_rules": expect_eligible,
                   "needs_followup": needs_followup, "notes": notes},
    }


# ================================================================
# EMAILS (50)
# ================================================================
print("Generating 50 emails...")
dup_sources = {}

for i in range(1, 51):
    sid = f"E{i:02d}"
    amount = random.choice([500, 800, 1200, 1500, 2000, 2500, 3000, 5000])
    date = future_date()
    files = []

    if i <= 15:      # digital PDF attachment, in-region, complete
        org = mk_org(i); contact = mk_contact(i, org["name"], freemail=(i % 5 == 0))
        pdf = os.path.join(DIR_EMAILS, f"{sid}_anfrage.pdf")
        write_pdf(pdf, letter_lines(org, contact, amount, date))
        with open(pdf, "rb") as f: data = f.read()
        eml = os.path.join(DIR_EMAILS, f"{sid}.eml")
        write_eml(eml, contact["email"], f"Sponsoringanfrage {org['name']}",
                  email_body(org, contact), [(f"{sid}_anfrage.pdf", data, "application", "pdf")])
        files = [f"{sid}.eml", f"{sid}_anfrage.pdf"]
        if i == 1: dup_sources["E01"] = (org, contact, data)
        record(sid, "email", files, gt(org, contact, amount, date, notes="clean digital PDF"))

    elif i <= 20:    # DOCX attachment
        org = mk_org(i); contact = mk_contact(i, org["name"])
        dx = os.path.join(DIR_EMAILS, f"{sid}_anfrage.docx")
        actual = write_docx(dx, letter_lines(org, contact, amount, date))
        with open(actual, "rb") as f: data = f.read()
        fname = os.path.basename(actual)
        write_eml(os.path.join(DIR_EMAILS, f"{sid}.eml"), contact["email"],
                  f"Anfrage Sponsoring - {org['name']}", email_body(org, contact),
                  [(fname, data, "application", "octet-stream")])
        if i == 16: dup_sources["E16"] = (org, contact, data, fname)
        record(sid, "email", [f"{sid}.eml", fname], gt(org, contact, amount, date, notes="DOCX attachment"))

    elif i <= 25:    # scanned-image attachment -> OCR path
        org = mk_org(i); contact = mk_contact(i, org["name"], freemail=True)
        jpg = os.path.join(DIR_EMAILS, f"{sid}_scan.jpg")
        write_scan_jpg(jpg, letter_lines(org, contact, amount, date))
        with open(jpg, "rb") as f: data = f.read()
        write_eml(os.path.join(DIR_EMAILS, f"{sid}.eml"), contact["email"],
                  f"Sponsoring {org['name']} (Scan)", email_body(org, contact),
                  [(f"{sid}_scan.jpg", data, "image", "jpeg")])
        record(sid, "email", [f"{sid}.eml", f"{sid}_scan.jpg"],
               gt(org, contact, amount, date, notes="scanned image -> OCR/vision path"))

    elif i <= 30:    # body-only, no attachment
        org = mk_org(i); contact = mk_contact(i, org["name"])
        body = "\n".join(letter_lines(org, contact, amount, date))
        write_eml(os.path.join(DIR_EMAILS, f"{sid}.eml"), contact["email"],
                  f"Sponsoringanfrage {org['name']}", body, [])
        record(sid, "email", [f"{sid}.eml"], gt(org, contact, amount, date, notes="body-only email"))

    elif i <= 36:    # incomplete -> follow-up loop
        org = mk_org(i); contact = mk_contact(i, org["name"])
        missing = [["amount"], ["date"], ["amount", "date"]][(i - 31) % 3]
        pdf = os.path.join(DIR_EMAILS, f"{sid}_anfrage.pdf")
        write_pdf(pdf, letter_lines(org, contact, amount, date, missing=missing))
        with open(pdf, "rb") as f: data = f.read()
        write_eml(os.path.join(DIR_EMAILS, f"{sid}.eml"), contact["email"],
                  f"Unterstuetzung fuer {org['name']}", email_body(org, contact),
                  [(f"{sid}_anfrage.pdf", data, "application", "pdf")])
        record(sid, "email", [f"{sid}.eml", f"{sid}_anfrage.pdf"],
               gt(org, contact, None if "amount" in missing else amount,
                  None if "date" in missing else date, complete=False, missing=missing,
                  needs_followup=True, notes=f"incomplete: missing {missing}"))

    elif i <= 41:    # out-of-region
        town = OUT_REGION_TOWNS[(i - 37) % len(OUT_REGION_TOWNS)]
        org = mk_org(i, town=town); contact = mk_contact(i, org["name"])
        pdf = os.path.join(DIR_EMAILS, f"{sid}_anfrage.pdf")
        write_pdf(pdf, letter_lines(org, contact, amount, date))
        with open(pdf, "rb") as f: data = f.read()
        write_eml(os.path.join(DIR_EMAILS, f"{sid}.eml"), contact["email"],
                  f"Sponsoringanfrage {org['name']}", email_body(org, contact),
                  [(f"{sid}_anfrage.pdf", data, "application", "pdf")])
        record(sid, "email", [f"{sid}.eml", f"{sid}_anfrage.pdf"],
               gt(org, contact, amount, date, in_region=False,
                  notes=f"out-of-region ({town}) - B29 region-handling test"))

    elif i <= 44:    # junk / not a sponsorship request
        kind = ["newsletter", "invoice", "job_application"][(i - 42)]
        bodies = {
            "newsletter": ("Newsletter Energiewende KW28",
                           "Unser woechentlicher Newsletter: Tipps zum Energiesparen im Sommer..."),
            "invoice": ("Rechnung Nr. 2026-4711",
                        "Sehr geehrte Damen und Herren,\n\nanbei unsere Rechnung ueber 238,00 EUR "
                        "fuer Bueromaterial. Zahlbar innerhalb 14 Tagen.\n\nBuchhaltung Mueller GmbH"),
            "job_application": ("Bewerbung als Elektroniker",
                                "Sehr geehrte Damen und Herren,\n\nhiermit bewerbe ich mich um die "
                                "ausgeschriebene Stelle als Elektroniker.\n\nMit freundlichen Gruessen\nKlaus Kandidat"),
        }
        subj, body = bodies[kind]
        write_eml(os.path.join(DIR_EMAILS, f"{sid}.eml"),
                  f"noreply{i}@{random.choice(FREEMAIL)}", subj, body, [])
        record(sid, "email", [f"{sid}.eml"],
               gt(None, None, None, None, in_region=False, complete=False,
                  is_sponsorship=False, expect_eligible=False,
                  notes=f"junk: {kind} -> classifier/relevance-gate must stop this"))

    elif i <= 46:    # duplicates of E01 / E16
        src = "E01" if i == 45 else "E16"
        if src == "E01":
            org, contact, data = dup_sources["E01"]; fname = f"{src}_anfrage.pdf"; sub = "pdf"
            atts = [(fname, data, "application", "pdf")]
        else:
            org, contact, data, fname = dup_sources["E16"]
            atts = [(fname, data, "application", "octet-stream")]
        other = mk_contact(i + 7, org["name"], freemail=True)
        write_eml(os.path.join(DIR_EMAILS, f"{sid}.eml"), other["email"],
                  f"WG: Sponsoringanfrage {org['name']}", "Zur Info, hatte ich schon geschickt?",
                  atts)
        record(sid, "email", [f"{sid}.eml"],
               gt(org, contact, None, None, complete=True,
                  notes=f"DUPLICATE of {src} (same attachment bytes, different sender) -> dedup test"))

    elif i == 47 or i == 48:  # multi-attachment
        org = mk_org(i); contact = mk_contact(i, org["name"])
        pdf = os.path.join(DIR_EMAILS, f"{sid}_anfrage.pdf")
        write_pdf(pdf, letter_lines(org, contact, amount, date))
        logo = os.path.join(DIR_EMAILS, f"{sid}_logo.jpg")
        pm = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 180, 100), False)
        pm.set_rect(pm.irect, (random.randint(0, 255), random.randint(0, 255), 200))
        pm.save(logo, output="jpeg")
        with open(pdf, "rb") as f: d1 = f.read()
        with open(logo, "rb") as f: d2 = f.read()
        write_eml(os.path.join(DIR_EMAILS, f"{sid}.eml"), contact["email"],
                  f"Sponsoringanfrage {org['name']} inkl. Logo", email_body(org, contact),
                  [(f"{sid}_anfrage.pdf", d1, "application", "pdf"),
                   (f"{sid}_logo.jpg", d2, "image", "jpeg")])
        record(sid, "email", [f"{sid}.eml", f"{sid}_anfrage.pdf", f"{sid}_logo.jpg"],
               gt(org, contact, amount, date, notes="multi-attachment (PDF + logo)"))

    elif i == 49:    # over-budget amount -> hard rule fail expected
        org = mk_org(i); contact = mk_contact(i, org["name"])
        amount = 60000
        pdf = os.path.join(DIR_EMAILS, f"{sid}_anfrage.pdf")
        write_pdf(pdf, letter_lines(org, contact, amount, date))
        with open(pdf, "rb") as f: data = f.read()
        write_eml(os.path.join(DIR_EMAILS, f"{sid}.eml"), contact["email"],
                  f"Grosssponsoring {org['name']}", email_body(org, contact),
                  [(f"{sid}_anfrage.pdf", data, "application", "pdf")])
        record(sid, "email", [f"{sid}.eml", f"{sid}_anfrage.pdf"],
               gt(org, contact, amount, date, expect_eligible=False,
                  notes="60000 EUR -> amount_range hard rule should fail"))

    else:            # i == 50: past event date
        org = mk_org(i); contact = mk_contact(i, org["name"])
        date = "2026-01-15"  # in the past relative to mid-2026
        pdf = os.path.join(DIR_EMAILS, f"{sid}_anfrage.pdf")
        write_pdf(pdf, letter_lines(org, contact, amount, date))
        with open(pdf, "rb") as f: data = f.read()
        write_eml(os.path.join(DIR_EMAILS, f"{sid}.eml"), contact["email"],
                  f"Sponsoringanfrage {org['name']}", email_body(org, contact),
                  [(f"{sid}_anfrage.pdf", data, "application", "pdf")])
        record(sid, "email", [f"{sid}.eml", f"{sid}_anfrage.pdf"],
               gt(org, contact, amount, date, expect_eligible=True,
                  notes="past event date -> event_date_future soft rule warning"))

# ================================================================
# WEB FORMS (50)
# ================================================================
print("Generating 50 web form payloads...")

for i in range(1, 51):
    sid = f"W{i:02d}"
    org = mk_org(i + 100)
    contact = mk_contact(i + 100, org["name"], freemail=(i % 6 == 0))
    amount = random.choice([400, 750, 1000, 1500, 2200, 3000, 4500])
    date = future_date()
    valid, notes, expect_eligible, in_region, needs_followup = True, "", True, True, False
    missing = []

    payload = {
        "organization_name": org["name"],
        "contact_name": contact["name"],
        "contact_email": contact["email"],
        "contact_phone": contact["phone"],
        "requested_amount": amount,
        "purpose": org["purpose"],
        "description": f"{org['purpose']} in {org['town']} - organisiert von {org['name']}.",
        "event_date": date,
        "target_audience": random.choice(["Familien", "Kinder und Jugendliche", "Senioren", "Alle Altersgruppen"]),
        "proposed_visibility": "Logo auf Plakaten, Website und Programmheft",
        "region": org["town"],
    }

    if i <= 35:
        notes = "valid complete in-region"
    elif i <= 40:   # schema-valid but incomplete -> completeness/follow-up
        drop = [["requested_amount"], ["event_date"], ["requested_amount", "event_date"]][(i - 36) % 3]
        for k in drop:
            payload[k] = None
        missing = drop; needs_followup = True
        amount = None if "requested_amount" in drop else amount
        date = None if "event_date" in drop else date
        notes = f"valid but incomplete: {drop}"
    elif i <= 44:   # out-of-region
        town = OUT_REGION_TOWNS[(i - 41) % len(OUT_REGION_TOWNS)]
        org = mk_org(i + 100, town=town)
        payload["organization_name"] = org["name"]
        payload["region"] = town
        payload["description"] = f"{org['purpose']} in {town}."
        in_region = False
        notes = f"out-of-region ({town})"
    elif i == 45:   # INVALID: missing required organization_name
        del payload["organization_name"]
        valid = False; notes = "INVALID: organization_name missing -> expect 422"
    elif i == 46:   # INVALID: bad email
        payload["contact_email"] = "keine-echte-adresse"
        valid = False; notes = "INVALID: bad contact_email -> expect 422"
    elif i == 47:   # INVALID: missing required purpose
        del payload["purpose"]
        valid = False; notes = "INVALID: purpose missing -> expect 422"
    elif i == 48:   # over budget
        amount = 75000; payload["requested_amount"] = amount
        expect_eligible = False; notes = "75000 EUR -> amount_range fail"
    elif i == 49:   # past date
        date = "2026-02-01"; payload["event_date"] = date
        notes = "past event date -> soft rule warning"
    else:           # i == 50: tiny amount
        amount = 50; payload["requested_amount"] = amount
        expect_eligible = False; notes = "50 EUR -> below minimum (if amount_range has a floor)"

    with open(os.path.join(DIR_WEBFORM, f"{sid}.json"), "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, ensure_ascii=False)

    files = [f"{sid}.json"]
    if valid and i <= 44:  # attachment for post-B01 testing
        att = os.path.join(DIR_WF_ATT, f"{sid}_anlage.pdf")
        write_pdf(att, letter_lines(org, contact, amount or 0, date or "2026-10-01",
                                    missing=missing))
        files.append(f"attachments/{sid}_anlage.pdf")

    record(sid, "web_form", files,
           gt(org if valid else None, contact if valid else None, amount, date,
              in_region=in_region, complete=(not missing and valid), missing=missing,
              is_sponsorship=valid, expect_eligible=expect_eligible and valid,
              needs_followup=needs_followup, notes=notes),
           payload=payload)

# ================================================================
# OPERATOR SCENARIOS (10)
# ================================================================
print("Generating 10 operator scenarios...")

OP_SCENARIOS = [
    ("O01", "fields + matching PDF", dict(fields=True, files=["pdf"], conflict=False)),
    ("O02", "fields CONFLICT with attached PDF (ground-truth test)", dict(fields=True, files=["pdf_conflict"], conflict=True)),
    ("O03", "fields only, no attachment", dict(fields=True, files=[])),
    ("O04", "file only, no fields", dict(fields=False, files=["pdf"])),
    ("O05", "NO contact email (dead-end/B24 test)", dict(fields=True, files=["pdf"], no_email=True)),
    ("O06", "fields + xlsx attachment (B04 test)", dict(fields=True, files=["xlsx"])),
    ("O07", "multi-file: PDF + scan JPG + xlsx", dict(fields=True, files=["pdf", "jpg", "xlsx"])),
    ("O08", "scan image only + minimal fields (OCR/vision)", dict(fields="minimal", files=["jpg"])),
    ("O09", "autopilot mode, complete", dict(fields=True, files=["pdf"], mode="autopilot")),
    ("O10", "incomplete fields + email (follow-up test)", dict(fields="incomplete", files=[])),
]

for idx, (sid, desc, cfg) in enumerate(OP_SCENARIOS):
    org = mk_org(idx + 200)
    contact = mk_contact(idx + 200, org["name"])
    amount = random.choice([600, 900, 1400, 1800, 2400])
    date = future_date()
    files = []

    scenario = {"id": sid, "description": desc,
                "pipeline_mode": cfg.get("mode", "copilot"), "fields": {}}

    if cfg["fields"]:
        f = {"contact_name": contact["name"],
             "contact_email": "" if cfg.get("no_email") else contact["email"],
             "contact_phone": contact["phone"],
             "organization_name": org["name"],
             "requested_amount": str(amount),
             "purpose": org["purpose"],
             "event_date": date, "region": org["town"]}
        if cfg["fields"] == "minimal":
            f = {"contact_name": contact["name"], "contact_email": contact["email"],
                 "contact_phone": "", "organization_name": "", "requested_amount": "",
                 "purpose": "", "event_date": "", "region": ""}
        if cfg["fields"] == "incomplete":
            f["requested_amount"] = ""; f["event_date"] = ""
        scenario["fields"] = f

    for kind in cfg["files"]:
        if kind == "pdf":
            p = os.path.join(DIR_OPERATOR, f"{sid}_anfrage.pdf")
            write_pdf(p, letter_lines(org, contact, amount, date))
            files.append(f"{sid}_anfrage.pdf")
        elif kind == "pdf_conflict":
            other = mk_org(idx + 250)   # different org entirely
            oc = mk_contact(idx + 250, other["name"])
            p = os.path.join(DIR_OPERATOR, f"{sid}_anfrage_konflikt.pdf")
            write_pdf(p, letter_lines(other, oc, amount + 1000, future_date()))
            files.append(f"{sid}_anfrage_konflikt.pdf")
        elif kind == "jpg":
            p = os.path.join(DIR_OPERATOR, f"{sid}_scan.jpg")
            write_scan_jpg(p, letter_lines(org, contact, amount, date))
            files.append(f"{sid}_scan.jpg")
        elif kind == "xlsx":
            try:
                import openpyxl
                wb = openpyxl.Workbook(); ws = wb.active
                ws.append(["Feld", "Wert"])
                ws.append(["Verein", org["name"]])
                ws.append(["Betrag", amount])
                ws.append(["Zweck", org["purpose"]])
                ws.append(["Datum", date])
                p = os.path.join(DIR_OPERATOR, f"{sid}_daten.xlsx")
                wb.save(p)
                files.append(f"{sid}_daten.xlsx")
            except ImportError:
                pass

    scenario["files"] = files
    with open(os.path.join(DIR_OPERATOR, f"{sid}.json"), "w", encoding="utf-8") as f:
        json.dump(scenario, f, indent=2, ensure_ascii=False)

    record(sid, "operator", [f"{sid}.json"] + files,
           gt(org, contact if not cfg.get("no_email") else {"email": None, **contact},
              amount if cfg["fields"] != "minimal" else None, date,
              complete=cfg["fields"] is True, is_sponsorship=True,
              needs_followup=cfg["fields"] in ("minimal", "incomplete"),
              notes=desc))

# ================================================================
# Manifest
# ================================================================
with open(os.path.join(OUT, "manifest.json"), "w", encoding="utf-8") as f:
    json.dump({"generated": "2026-07-08", "seed": 42, "total": len(MANIFEST),
               "region_config": "Stadtwerke Bodensee (in-region = Bodensee towns)",
               "samples": MANIFEST}, f, indent=2, ensure_ascii=False)

print(f"\nDONE: {len(MANIFEST)} samples in {OUT}")
print(f"  emails:   {sum(1 for m in MANIFEST if m['channel'] == 'email')}")
print(f"  webform:  {sum(1 for m in MANIFEST if m['channel'] == 'web_form')}")
print(f"  operator: {sum(1 for m in MANIFEST if m['channel'] == 'operator')}")
sys.exit(0)
