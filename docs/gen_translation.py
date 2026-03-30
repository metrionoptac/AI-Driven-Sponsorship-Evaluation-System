"""Generate the English translation DOCX for Bewertungskriterien."""
from docx import Document

doc = Document()
doc.add_heading("Bewertungskriterien - English Translation", level=0)
doc.add_paragraph("Source: Laura/Conoscope | Translated from German")

sections = {
    "1. Formal Eligibility / Exclusion": [
        "No political party", "No party-affiliated foundation",
        "No political actors or candidates", "No individuals (private persons)",
        "No profit-oriented organization", "No commercial mandatory task",
        "No purely commercial purpose", "No citizens' initiative",
        "No company sports group", "No homeowners' association",
        "No religious or church-affiliated organizations (if excluded)",
        "No relationship to employees / no conflict of interest",
        "No discrimination", "No glorification of violence",
        "No contradiction to corporate identity / mission statement",
        "No double funding with another funding format",
        "Within business territory / defined funding area / no foreign",
    ],
    "2. Organization and Master Data": [
        "Name of organization", "Partner / institution", "Type of organization",
        "Industry", "Category", "Division", "Sponsorship or donation",
        "Genre / topic area", "Project / event name", "Duration", "Year",
        "Contract / intended use", "Type of support", "Branch / location",
        "District / site", "Address", "Contact person", "Email / contact details",
        "Website", "Volunteer-run yes/no",
    ],
    "3. Request Quality / Completeness": [
        "All relevant basic information present", "Clear description of the project",
        "Comprehensible intended use", "Specific naming of desired support",
        "Clear timeframe", "Named target audience", "Budget / funding amount stated",
        "Structured and understandable presentation", "Missing information identifiable",
        "Need for follow-up questions", "Comparability with other requests",
        "Quality / level of detail of the request",
    ],
    "4. Regionality / Spatial Reference": [
        "Headquarters in relevant funding area", "Project location in relevant funding area",
        "District reference", "City area reference", "Regional reference",
        "Direct surroundings / local environment", "Supra-regional visibility",
        "National visibility", "Regional rootedness", "Local benefit for population / community",
    ],
    "5. Target Audience Reference": [
        "Primary target audience", "Target audience penetration", "Reaching broad audiences",
        "Reaching niche audiences", "Population in business territory",
        "Residents of a city / region", "Families", "Children and youth",
        "Elderly / seniors", "Disadvantaged groups", "Customers",
        "Employees / potential employees", "Business partners / network actors",
        "Multipliers / opinion leaders",
    ],
    "6. Community / Cohesion / Social Proximity": [
        "Community and partnership", "Proximity to people / citizen closeness",
        "Promoting togetherness", "Social cohesion", "Development of community feeling",
        "Reduction of prejudices", "Neighborhood reference", "Regional identification",
        "Local connectedness", "Improving quality of life",
    ],
    "7. Social / Societal Responsibility": [
        "Responsibility for society", "Social impact",
        "Support for socially disadvantaged groups", "Poverty reduction",
        "Help for homeless", "Debt counseling", "Support for those in need",
        "Social justice", "Equal opportunity",
        "Integration and anti-discrimination work",
        "Common good contribution", "Societal participation",
    ],
    "8. Education / Development / Participation": [
        "Education and change", "Quality education", "Reading promotion",
        "Tutoring programs", "Support for school equipment",
        "Education programs for disadvantaged students", "Education infrastructure",
        "Digital learning spaces", "Lifelong learning",
        "Mentoring programs", "Future perspectives for next generations",
    ],
    "9. Health / Prevention / Well-being": [
        "Health and well-being", "Health prevention", "Prevention programs",
        "Movement promotion", "Mental health", "Healthy nutrition",
        "Health checks", "Community-based health offerings",
    ],
    "10. Equality / Diversity / Inclusion": [
        "Gender equality", "Promotion of women and girls", "Diversity programs",
        "Inclusive workplace design", "Promotion of marginalized groups",
        "Anti-discrimination programs", "Integration of refugees",
        "Inclusion of people with disabilities",
    ],
    "11. Environment / Sustainability / Climate": [
        "Preserve environment and protect climate", "Ecological impact",
        "Climate protection", "CO2 reduction", "Energy efficiency",
        "Renewable energy", "Sustainable mobility", "Urban greening",
        "Biodiversity", "Water protection", "Recycling", "Environmental education",
    ],
    "12. Water / Supply / Public Services": [
        "Clean water and sanitation", "Water supply", "Drinking water fountains",
        "Public services (Daseinsvorsorge)", "Basic energy / water supply",
    ],
    "13. Economy / Work / Qualification": [
        "Decent work and economic growth", "Creation of training and jobs",
        "Workforce qualification", "Fair working conditions",
        "Support for regional economy", "Support for startups",
    ],
    "14. Innovation / Future / Infrastructure": [
        "Innovation and future", "Shaping the future", "Innovation-oriented projects",
        "Digitalization in schools or communities", "Hackathons",
        "Modern and sustainable infrastructure",
    ],
    "15. City / District / Living Space": [
        "Sustainable cities and communities", "Social and ecological urban development",
        "Livable place for all", "District associations", "Neighborhood projects",
        "Public transport promotion", "Local district development",
    ],
    "16. Culture / Homeland / Public Life": [
        "Culture and homeland", "Promotion of art and culture",
        "Events / festivals / markets", "Cultural education",
        "Public meeting character", "Tradition / homeland reference",
        "Strengthening local identity",
    ],
    "17. Sport / Movement / Club Life": [
        "Sport", "Amateur and recreational sport", "Professional sport",
        "Movement and prevention", "Youth work in sport", "Club life",
        "Local sport reference",
    ],
    "18. Image / Reputation / Perception": [
        "Positive image transfer", "Sympathy", "Down-to-earth quality", "Trust",
        "Credible connection between company and partner", "Reputation potential",
        "Perception as responsible actor", "Image score",
    ],
    "19. Visibility / Reach / Presence": [
        "Reach", "Regionality of reach", "Brand awareness", "Visibility general",
        "On-site visibility", "Media presence", "Social media reach",
        "Logo placement / advertising materials", "Direct target audience reach",
    ],
    "20. Activation / Communication / Storytelling": [
        "Activation potential", "Joint communication", "Storytelling",
        "Campaign capability", "Personal encounters", "Dialogue with citizens",
        "Communication occasions", "Interaction opportunities",
        "Value beyond mere logo placement",
    ],
    "21. Relationship / Network / Door Opener": [
        "Network", "Door-opener potential", "Access to relevant target audiences",
        "Access to communities", "Access to decision-makers", "New partnerships",
        "Public-private partnerships", "Cooperations",
    ],
    "22. Customer / Market / Business Benefit": [
        "Customer retention", "Sales potential", "Key account relevance",
        "Employer brand", "Company positioning", "Market development",
        "Strategic impact", "Impact score",
    ],
    "23. Portfolio / Strategy / Priority Setting": [
        "Portfolio fit", "Genre fit", "Strategic context", "Fit with priority areas",
        "Complementing existing portfolio", "Avoiding redundancies",
        "Balance of topic areas",
    ],
    "24. Budget / Capacity / Feasibility": [
        "Budget availability", "Current sponsorship budget", "Personnel capacity",
        "Organizational feasibility", "Operational effort", "Timing fit",
        "Effort-benefit ratio",
    ],
    "25. Sustainability Goals / SDG Reference": [
        "No poverty", "No hunger", "Health and well-being", "Quality education",
        "Gender equality", "Clean water and sanitation", "Affordable and clean energy",
        "Decent work and economic growth", "Industry innovation infrastructure",
        "Reduced inequalities", "Sustainable cities", "Responsible consumption",
        "Climate action", "Life below water", "Life on land",
        "Peace justice strong institutions", "Partnerships for goals",
    ],
    "26. Decision / Reasoning / Result Logic": [
        "Formal status", "Fulfilled / open / review", "Category suggestion",
        "Overall value", "Clear reasoning for approval", "Clear reasoning for rejection",
        "Clear reasoning for follow-up", "Standardized follow-up questions",
        "Transparent decision", "Traceable decision", "Documentable result",
        "Consistent reasoning across comparable cases",
    ],
}

for title, items in sections.items():
    doc.add_heading(title, level=2)
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

doc.save("docs/Bewertungskriterien_EN_Translation.docx")
print("DONE")
