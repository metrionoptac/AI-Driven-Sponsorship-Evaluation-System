"""
DOCX document parser.
Extracts text from .docx files preserving structure (paragraphs, tables, metadata).
"""

import logging
from dataclasses import dataclass

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)


@dataclass
class DOCXResult:
    """Result of DOCX extraction."""
    text: str
    tables: list[list[list[str]]]     # list of tables, each table = list of rows, each row = list of cells
    metadata: dict
    confidence: float = 1.0
    method: str = "python-docx"


def extract_docx(docx_bytes: bytes) -> DOCXResult:
    """
    Extract text from .docx bytes.

    Extracts:
    - All paragraphs (in document order)
    - All tables (as structured data)
    - Metadata (author, created date, etc.)
    """
    import io

    try:
        doc = Document(io.BytesIO(docx_bytes))
    except (PackageNotFoundError, Exception) as e:
        logger.error("Failed to open DOCX: %s", e)
        return DOCXResult(text="", tables=[], metadata={}, confidence=0.0)

    # Extract metadata
    props = doc.core_properties
    metadata = {
        "author": props.author or "",
        "title": props.title or "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
        "subject": props.subject or "",
    }

    # Extract paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style and para.style.name and "Heading" in para.style.name:
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    level_num = int(level)
                    text = f"{'#' * level_num} {text}"
                except ValueError:
                    text = f"## {text}"
            paragraphs.append(text)

    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            tables.append(table_data)

    # Combine paragraphs and format tables
    text_parts = list(paragraphs)

    for i, table in enumerate(tables):
        if table:
            # Format table as readable text
            text_parts.append(f"\n[Table {i + 1}]")
            for row in table:
                text_parts.append(" | ".join(row))

    full_text = "\n\n".join(text_parts)

    logger.info(
        "DOCX extracted: %d paragraphs, %d tables, %d chars",
        len(paragraphs), len(tables), len(full_text),
    )

    return DOCXResult(
        text=full_text,
        tables=tables,
        metadata=metadata,
        confidence=1.0,
    )
